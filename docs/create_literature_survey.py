import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ASCENDING_OUTPUT = Path(__file__).with_name("verisight_literature_survey_2022_mid_2026.docx")
DESCENDING_OUTPUT = Path(__file__).with_name("verisight_literature_survey_2026_to_2022.docx")
CORE_OUTPUT = Path(__file__).with_name("verisight_core_literature_survey_18_papers.docx")


# Curated, project-relevant peer-reviewed literature. "et al." is used where
# a paper has a long author list; the linked official publisher page lists all authors.
PAPERS = [
    ("2022", "TRUE: Re-evaluating Factual Consistency Evaluation", "Or Honovich et al.", "NAACL-HLT", "Metrics were tested in isolated tasks, so example-level factuality was unclear.", "Standardizes evaluation across 11 datasets and compares NLI with QA-based metrics.", "Supports VeriSight's NLI verifier and benchmark-based evaluation.", "https://aclanthology.org/2022.naacl-main.287/"),
    ("2022", "QAFactEval: Improved QA-Based Factual Consistency Evaluation for Summarization", "Alexander Fabbri et al.", "NAACL-HLT", "Existing factuality metrics gave inconsistent results across setups.", "Uses generated questions and answers to compare output with source material.", "A candidate future QA-based verifier to complement NLI.", "https://aclanthology.org/2022.naacl-main.187/"),
    ("2023", "SelfCheckGPT: Zero-Resource Black-Box Hallucination Detection", "Potsawee Manakul et al.", "EMNLP", "Closed LLMs hide internal probabilities and may not have external knowledge access.", "Samples multiple responses and flags factual disagreement.", "Foundation for VeriSight's self-consistency uncertainty score.", "https://aclanthology.org/2023.emnlp-main.557/"),
    ("2023", "HaluEval: A Large-Scale Hallucination Evaluation Benchmark", "Junyi Li et al.", "EMNLP", "Few large, human-annotated hallucination benchmarks existed for LLMs.", "Provides factual and hallucinated examples for recognition evaluation.", "Used by VeriSight to measure accuracy, Macro-F1, risk precision, recall, F1 and latency.", "https://aclanthology.org/2023.emnlp-main.397/"),
    ("2023", "FActScore: Fine-grained Atomic Evaluation of Factual Precision", "Sewon Min et al.", "EMNLP", "A whole answer may mix supported and unsupported facts; one verdict hides this.", "Splits text into atomic facts and measures the share supported by reliable knowledge.", "Direct basis for claim extraction, per-claim verdicts and reliability scoring.", "https://aclanthology.org/2023.emnlp-main.741/"),
    ("2023", "Hallucination Detection for Generative LLMs by Bayesian Sequential Estimation", "Xiaohua Wang et al.", "EMNLP", "Retrieving much evidence for every answer increases latency and cost.", "Uses Bayesian stop-or-continue decisions to reduce unnecessary evidence checks.", "Motivates latency-aware verification and early stopping in future versions.", "https://aclanthology.org/2023.emnlp-main.949/"),
    ("2023", "A New Benchmark and Reverse Validation Method for Passage-level Hallucination Detection", "Shiping Yang et al.", "Findings of EMNLP", "Earlier zero-resource methods focused mainly on individual sentences.", "Introduces reverse validation and a human-annotated passage-level benchmark.", "Relevant to VeriSight's long answers and future paragraph-level analysis.", "https://aclanthology.org/2023.findings-emnlp.256/"),
    ("2023", "Enhancing Uncertainty-Based Hallucination Detection with Stronger Focus", "Tianhang Zhang et al.", "EMNLP", "Retrieval and multi-sample consistency checks can be expensive.", "Uses token-level uncertainty cues without external references.", "Supports the reasoning behind an uncertainty indicator separate from reliability.", "https://aclanthology.org/2023.emnlp-main.58/"),
    ("2024", "Self-RAG: Learning to Retrieve, Generate and Critique through Self-Reflection", "Akari Asai et al.", "ICLR", "Fixed retrieval can add irrelevant passages or be used when not needed.", "Lets a model retrieve on demand and critique its generation with reflection tokens.", "Supports VeriSight's retrieve-generate-verify-correct design.", "https://openreview.net/forum?id=hSyW5go0v8"),
    ("2024", "ARES: An Automated Evaluation Framework for Retrieval-Augmented Generation Systems", "Jon Saad-Falcon et al.", "NAACL-HLT", "RAG evaluation normally needs costly human labels for every query, context and answer.", "Creates synthetic data and evaluates context relevance, faithfulness and answer relevance.", "Inspires separate evidence quality, answer reliability and evaluation metrics.", "https://aclanthology.org/2024.naacl-long.20/"),
    ("2024", "RAGAs: Automated Evaluation of Retrieval Augmented Generation", "Shahul Es et al.", "EACL Demo", "RAG evaluation often lacks reference answers and human labels.", "Introduces automated measures for context relevance, faithfulness and answer relevance.", "Related to VeriSight's evidence quality and source-agreement reporting.", "https://aclanthology.org/2024.eacl-demo.16/"),
    ("2024", "RAGTruth: A Hallucination Corpus for Developing Trustworthy RAG", "Cheng Niu et al.", "ACL", "RAG can still generate unsupported or contradictory claims; dedicated labels were limited.", "Provides nearly 18,000 manually annotated RAG responses with case- and word-level labels.", "Strong future benchmark for document/web-grounded claims and span highlighting.", "https://aclanthology.org/2024.acl-long.585/"),
    ("2024", "Knowledge-Centric Hallucination Detection", "Xiangkun Hu et al.", "EMNLP", "Response- or sentence-level checking can be too coarse for factual errors.", "RefChecker extracts claim triplets and checks them against a reference.", "Very close to VeriSight's fine-grained claim/evidence verification pipeline.", "https://aclanthology.org/2024.emnlp-main.395/"),
    ("2024", "Unsupervised Real-Time Hallucination Detection based on Internal States of LLMs", "Weihang Su et al.", "Findings of ACL", "Post-processing detectors can be slow and disconnected from generation.", "MIND uses model internal states for real-time detection without manual labels.", "A future local/open-model alternative; VeriSight presently uses external evidence instead.", "https://aclanthology.org/2024.findings-acl.854/"),
    ("2024", "Two-tiered Encoder-based Hallucination Detection for RAG in the Wild", "Ilana Zimmerman et al.", "EMNLP Industry Track", "Public benchmarks often miss production latency, domain shifts and non-verifiable chat.", "Tests encoder detectors on enterprise data and combines public with in-domain training.", "Supports separating factual claims from non-factual requests in VeriSight.", "https://aclanthology.org/2024.emnlp-industry.2/"),
    ("2024", "RAGChecker: A Fine-grained Framework for Diagnosing RAG", "Dongyu Ru et al.", "NeurIPS Datasets & Benchmarks", "A single final score cannot identify whether retrieval or generation failed.", "Uses diagnostic metrics for retrieval and generation modules separately.", "Supports VeriSight's modular quality, agreement, citation and verdict outputs.", "https://arxiv.org/abs/2408.08067"),
    ("2024", "Detecting Hallucinations in LLMs Using Semantic Entropy", "Sebastian Farquhar et al.", "Nature", "Token-level uncertainty treats different wordings of the same idea as different answers.", "Calculates uncertainty at semantic meaning level across multiple generations.", "Theoretical support for reporting uncertainty independently of evidence reliability.", "https://www.nature.com/articles/s41586-024-07421-0"),
    ("2024", "Hallucinations in Large Language Models (LLMs)", "G. Pradeep Reddy et al.", "IEEE eStream", "Users can mistake fluent, plausible text for factual text.", "Reviews hallucination causes, impacts, detection and mitigation strategies.", "Useful IEEE background reference for the problem statement and motivation.", "https://ieeexplore.ieee.org/document/10542617"),
    ("2025", "Beyond Facts: Evaluating Intent Hallucination in Large Language Models", "Yijie Hao et al.", "ACL", "An answer can contain facts but still omit or misinterpret the user's requested constraints.", "Introduces FAITHQA and an intent-constraint metric.", "Relevant to improving follow-up questions and answer relevance.", "https://aclanthology.org/2025.acl-long.349/"),
    ("2025", "REFIND: Retrieval-Augmented Factuality Hallucination Detection", "DongGeon Lee, Hwanjo Yu", "SemEval", "Answer-level labels do not show which exact text is unsupported.", "Detects hallucinated spans using retrieved documents and context-sensitivity ratio.", "Relevant to highlighted claims/spans and evidence-sensitive checking.", "https://aclanthology.org/2025.semeval-1.2/"),
    ("2025", "Learning Auxiliary Tasks Improves Reference-Free Hallucination Detection in Open-Domain Long-Form Generation", "Chengwei Qin et al.", "ACL", "Internal confidence alone is not reliably better than random on open-domain factuality.", "Uses auxiliary tasks to improve reference-free detection.", "Explains why VeriSight should not rely only on LLM confidence; it retrieves evidence.", "https://aclanthology.org/2025.acl-short.93/"),
    ("2025", "Long-form Hallucination Detection with Self-elicitation", "Zihang Liu et al.", "Findings of ACL", "Isolated claim checks can miss relationships across a long response.", "Uses self-elicited knowledge and graphs to preserve contextual semantics.", "Relevant for future long-answer and multi-claim dependency checks.", "https://aclanthology.org/2025.findings-acl.211/"),
    ("2025", "Simple Factuality Probes Detect Hallucinations in Long-Form NLG", "Jiatong Han et al.", "Findings of EMNLP", "Multi-sample hallucination detection can be computationally expensive.", "Uses lightweight probes on hidden states with one generated sample.", "Future optimisation path for local/open models and lower verification cost.", "https://aclanthology.org/2025.findings-emnlp.880/"),
    ("2025", "RAC: Efficient LLM Factuality Correction with Retrieval Augmentation", "Changmao Li, Jeffrey Flanigan", "Findings of EMNLP", "Detection alone leaves users without a corrected, evidence-based answer.", "Decomposes facts, retrieves evidence, verifies and corrects with low latency.", "Closest published match to VeriSight's correction-with-citations workflow.", "https://aclanthology.org/2025.findings-emnlp.1370/"),
    ("2025", "RAGGED: Towards Informed Design of Scalable and Stable RAG Systems", "Jennifer Hsia et al.", "ICML", "RAG results vary with retriever settings, source count and system configuration.", "Defines metrics for RAG stability and scalability.", "Relevant when tuning source limits, reranking and response latency.", "https://proceedings.mlr.press/v267/hsia25a.html"),
    ("2025", "Alleviating Hallucinations of LLMs through Induced Hallucinations", "Yue Zhang et al.", "Findings of NAACL", "Fine-tuning or expensive external tools are not always available.", "Uses induce-then-contrast decoding to suppress untruthful token predictions.", "Future prevention-layer alternative; not required for VeriSight's post-hoc verifier.", "https://aclanthology.org/2025.findings-naacl.459/"),
    ("2025", "Removal of Hallucination on Hallucination: Debate-Augmented RAG", "Wentao Hu et al.", "ACL", "Bad retrieval can itself introduce a second source of hallucination.", "Uses multi-agent debate to refine both retrieval and generation.", "Supports checking evidence quality before trusting a citation or generated answer.", "https://aclanthology.org/2025.acl-long.770/"),
    ("2026", "RLSeek: Evidence-Grounded Reasoning for RAG Hallucination Detection", "Zhaoheng Huang et al.", "ACL", "Reasoning-based detectors may judge without explicitly grounding claims in retrieved text.", "Trains evidence-grounded reasoning that quotes/verifies against retrieved sources.", "Strong support for VeriSight's claim-to-evidence verification and citations.", "https://aclanthology.org/2026.acl-long.1492/"),
    ("2026", "ReFL: Reflective Feedback Learning for Hallucination Detection of LLMs", "Cunhang Fan et al.", "ACL", "External retrieval may be slow; hidden-state methods may generalize poorly.", "Uses corrective in-context feedback without updating model weights.", "A future enhancement for improved detection without model fine-tuning.", "https://aclanthology.org/2026.acl-long.899/"),
    ("2026", "MARCH: Multi-Agent Reinforced Check for Hallucination", "Zhuo Li et al.", "ACL", "LLM-as-a-judge can reproduce the original generator's errors through confirmation bias.", "Uses solver, proposer and checker agents with information asymmetry.", "Useful comparison point; VeriSight avoids relying on one LLM judge by using evidence and NLI.", "https://aclanthology.org/2026.acl-long.1828/"),
    ("2026", "HAD: HAllucination Detection Language Models Based on a Comprehensive Taxonomy", "Fan Xu et al.", "ACL Industry Track", "Narrow taxonomies and separate detectors do not cover detection, localization and correction together.", "Introduces 11 hallucination categories plus joint detection, span identification and correction.", "Supports VeriSight's detect-explain-correct scope and future taxonomy expansion.", "https://aclanthology.org/2026.acl-industry.11/"),
    ("2026", "Enhancing Hallucination Detection via Future Context", "Joosung Lee et al.", "Findings of ACL", "Black-box sampling methods use only current generated text and miss later warning signals.", "Samples future context because hallucinations tend to persist.", "A future improvement to VeriSight's uncertainty/self-consistency method.", "https://aclanthology.org/2026.findings-acl.35/"),
    ("2026", "PROBE: PROcess-Based BEnchmark for Hallucination Detection", "Yu Zhang et al.", "Findings of ACL", "One-shot LLM judging gives little transparency about where verification fails.", "Benchmarks four steps: claim decomposition, evidence finding, evidence evaluation and localization.", "Almost exactly matches VeriSight's modular pipeline and evaluation approach.", "https://aclanthology.org/2026.findings-acl.2099/"),
    ("2026", "Logical Consistency as a Bridge: Improving LLM Hallucination Detection", "Hao Mi et al.", "ACL", "Methods use either internal uncertainty or verbal self-judgment, but not both together.", "Combines implicit and explicit signals through mutual learning and label constraints.", "Future ensemble-verifier idea for combining uncertainty and evidence verdicts.", "https://aclanthology.org/2026.acl-long.286/"),
    ("2026", "Re3: Relevance & Recency Retrieval for Mitigating Temporal Hallucination", "Jiawei Cao et al.", "ACL", "Retrievers can return outdated information when facts change over time.", "Uses time-aware relevance and a conflict-aware recency filter.", "Relevant to current-affairs questions and source freshness checks.", "https://aclanthology.org/2026.acl-long.1180/"),
    ("2026", "Detecting Hallucinations in RAG via Semantic-level Internal Reasoning Graph", "Xiangkun Hu et al.", "Findings of ACL", "Sentence-level checking can miss semantic dependencies between evidence and claims.", "Builds an internal reasoning graph for semantic-level RAG verification.", "Future reference for graph-based evidence/claim linking in VeriSight.", "https://aclanthology.org/2026.findings-acl.1385/"),
]


# A concise set for the final report: every item supports a component that is
# already implemented in VeriSight, rather than a future-only enhancement.
CORE_TITLES = {
    "PROBE: PROcess-Based BEnchmark for Hallucination Detection",
    "RAC: Efficient LLM Factuality Correction with Retrieval Augmentation",
    "REFIND: Retrieval-Augmented Factuality Hallucination Detection",
    "Beyond Facts: Evaluating Intent Hallucination in Large Language Models",
    "RAGChecker: A Fine-grained Framework for Diagnosing RAG",
    "Knowledge-Centric Hallucination Detection",
    "ARES: An Automated Evaluation Framework for Retrieval-Augmented Generation Systems",
    "RAGAs: Automated Evaluation of Retrieval Augmented Generation",
    "RAGTruth: A Hallucination Corpus for Developing Trustworthy RAG",
    "Detecting Hallucinations in LLMs Using Semantic Entropy",
    "SelfCheckGPT: Zero-Resource Black-Box Hallucination Detection",
    "HaluEval: A Large-Scale Hallucination Evaluation Benchmark",
    "FActScore: Fine-grained Atomic Evaluation of Factual Precision",
    "Hallucination Detection for Generative LLMs by Bayesian Sequential Estimation",
    "TRUE: Re-evaluating Factual Consistency Evaluation",
    "QAFactEval: Improved QA-Based Factual Consistency Evaluation for Summarization",
}

EXTRA_CORE_PAPERS = [
    (
        "2025",
        "Trust Me, I'm Wrong: LLMs Hallucinate with Certainty Despite Knowing the Answer",
        "Adi Simhi et al.",
        "arXiv preprint",
        "Users can mistake a fluent answer or the model's apparent confidence for factual correctness.",
        "Shows that models can produce confident hallucinations even when they contain the correct knowledge.",
        "Justifies VeriSight's evidence-based reliability score instead of trusting a generator's confidence.",
        "https://arxiv.org/abs/2502.12964",
    ),
    (
        "2025",
        "Uncertainty Quantification for Language Models: A Suite of Black-Box, White-Box, LLM Judge, and Ensemble Scorers",
        "Dylan Bouchard, Mohit Singh Chauhan et al.",
        "TMLR",
        "Uncertainty methods are often compared inconsistently across different model-access settings.",
        "Systematically compares black-box, white-box, LLM-judge and ensemble uncertainty scorers.",
        "Supports VeriSight's black-box multi-response uncertainty estimate for Gemini and Groq.",
        "https://arxiv.org/abs/2504.19254",
    ),
]


def core_papers():
    selected = [paper for paper in PAPERS if paper[1] in CORE_TITLES]
    selected.extend(EXTRA_CORE_PAPERS)
    return sorted(selected, key=lambda paper: (int(paper[0]), paper[1]), reverse=True)


# Method and outcome are intentionally qualitative unless the source paper
# reports a directly comparable benchmark number. Each paper uses a different
# task or dataset, so their scores must not be ranked against one another.
CORE_METHOD_AND_RESULT = {
    "PROBE: PROcess-Based BEnchmark for Hallucination Detection": (
        "Process-level evaluation: claim decomposition, evidence finding, evidence evaluation and localisation.",
        "Provides a diagnostic benchmark rather than one universal best score; shows that failures can occur at any verification stage.",
    ),
    "Trust Me, I'm Wrong: LLMs Hallucinate with Certainty Despite Knowing the Answer": (
        "Confidence-versus-correctness analysis of LLM answers.",
        "Finds that apparent model confidence is not a reliable factuality signal; external evidence is needed.",
    ),
    "Uncertainty Quantification for Language Models: A Suite of Black-Box, White-Box, LLM Judge, and Ensemble Scorers": (
        "Comparison of black-box sampling, white-box, judge and ensemble uncertainty scorers.",
        "No universal winner across models; black-box multi-response methods are practical when internal probabilities are unavailable.",
    ),
    "RAC: Efficient LLM Factuality Correction with Retrieval Augmentation": (
        "Fact decomposition, retrieval, verification and evidence-grounded correction.",
        "Reports effective factuality correction with retrieval while targeting lower latency than repeated full-answer regeneration.",
    ),
    "REFIND: Retrieval-Augmented Factuality Hallucination Detection": (
        "Retrieved-document span detection using context-sensitivity signals.",
        "Outperforms span-detection baselines on its SemEval evaluation setting; supports localising unsupported text.",
    ),
    "Beyond Facts: Evaluating Intent Hallucination in Large Language Models": (
        "Intent-constraint evaluation with the FAITHQA benchmark.",
        "Shows that factual answers may still fail the user's requested intent or constraints.",
    ),
    "RAGChecker: A Fine-grained Framework for Diagnosing RAG": (
        "Fine-grained RAG diagnostics for retrieval and generation components.",
        "Provides more actionable diagnosis than a single end-to-end score by identifying the failing pipeline component.",
    ),
    "Knowledge-Centric Hallucination Detection": (
        "Claim-triplet extraction and evidence checking (RefChecker).",
        "Reports stronger factuality assessment than coarse answer-level checking across its reference-based evaluations.",
    ),
    "ARES: An Automated Evaluation Framework for Retrieval-Augmented Generation Systems": (
        "Synthetic-data-assisted RAG evaluation of context relevance, faithfulness and answer relevance.",
        "Produces automated RAG metrics designed to align with human evaluation without labelling every production example.",
    ),
    "RAGAs: Automated Evaluation of Retrieval Augmented Generation": (
        "Reference-free RAG measures for context relevance, faithfulness and answer relevance.",
        "Demonstrates practical RAG evaluation without requiring a gold reference answer for every query.",
    ),
    "RAGTruth: A Hallucination Corpus for Developing Trustworthy RAG": (
        "Human-annotated RAG hallucination corpus with response- and word-level labels.",
        "Provides nearly 18,000 labelled RAG responses for evaluating grounded hallucination detection.",
    ),
    "Detecting Hallucinations in LLMs Using Semantic Entropy": (
        "Semantic-meaning uncertainty over multiple generated responses.",
        "Shows that grouping answers by meaning can detect uncertainty better than comparing only individual tokens.",
    ),
    "SelfCheckGPT: Zero-Resource Black-Box Hallucination Detection": (
        "Self-consistency checking across multiple independently sampled answers.",
        "Shows that disagreement between samples is a useful hallucination signal without external model access.",
    ),
    "HaluEval: A Large-Scale Hallucination Evaluation Benchmark": (
        "Human-annotated hallucination-recognition benchmark.",
        "Supplies fixed grounded and hallucinated cases for comparable accuracy, Macro-F1, precision and recall measurement.",
    ),
    "FActScore: Fine-grained Atomic Evaluation of Factual Precision": (
        "Atomic-fact decomposition followed by support checking against knowledge.",
        "Shows why claim-level precision reveals mixed correct/incorrect answers that an answer-level score hides.",
    ),
    "Hallucination Detection for Generative LLMs by Bayesian Sequential Estimation": (
        "Bayesian sequential stop-or-continue evidence verification.",
        "Reduces unnecessary evidence checks while preserving a targeted hallucination-detection decision process.",
    ),
    "TRUE: Re-evaluating Factual Consistency Evaluation": (
        "Cross-task comparison of NLI- and QA-based factual-consistency metrics across 11 datasets.",
        "Establishes NLI as a strong, reusable factual-consistency baseline across multiple grounded-generation tasks.",
    ),
    "QAFactEval: Improved QA-Based Factual Consistency Evaluation for Summarization": (
        "Question generation and answering to compare a response with its source.",
        "Reports improved factual-consistency correlation over earlier QA-based evaluation baselines.",
    ),
}


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_hyperlink(paragraph, text, url, font_size=7.5, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(font_size * 2)))
    r_pr.append(size)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def write_cell(cell, text, font_size=7.5, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def main():
    parser = argparse.ArgumentParser(description="Build the VeriSight literature survey.")
    parser.add_argument(
        "--descending",
        action="store_true",
        help="Sort papers from 2026 down to 2022 and write the descending edition.",
    )
    parser.add_argument(
        "--core",
        action="store_true",
        help="Create the focused 18-paper survey for the implemented VeriSight pipeline.",
    )
    args = parser.parse_args()
    if args.core:
        papers = core_papers()
        output = CORE_OUTPUT
        order_label = "Core current-implementation papers | Descending order: 2026 to 2022"
    else:
        papers = sorted(PAPERS, key=lambda paper: int(paper[0]), reverse=args.descending)
        output = DESCENDING_OUTPUT if args.descending else ASCENDING_OUTPUT
        order_label = "Descending order: 2026 to 2022" if args.descending else "Ascending order: 2022 to 2026"

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.25 if args.core else 0.35)
    section.bottom_margin = Inches(0.25 if args.core else 0.35)
    section.left_margin = Inches(0.2 if args.core else 0.3)
    section.right_margin = Inches(0.2 if args.core else 0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("VeriSight Literature Survey: Hallucination Detection and Evidence-Based Verification")
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(31, 78, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    sr = subtitle.add_run(f"Curated, project-relevant literature from 2022 through July 2026 | {len(papers)} papers | {order_label} | Direct publisher links included")
    sr.italic = True
    sr.font.name = "Arial"
    sr._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    sr._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    sr.font.size = Pt(8.5)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    scope = (
        "Scope note: This focused review contains only papers that directly support VeriSight's currently implemented pipeline - answer generation, claim decomposition, evidence retrieval, NLI/evidence verification, reliability, uncertainty, correction, citations, and evaluation. "
        if args.core
        else "Scope note: This is a broad, curated review of papers directly supporting VeriSight's pipeline - retrieval, claim decomposition, NLI/evidence verification, uncertainty, correction, citations, and evaluation. "
    )
    nr = note.add_run(scope + "It is not a claim that every paper published worldwide on hallucination has been listed. URLs are displayed in full and embedded as Word hyperlinks for printing or opening later.")
    nr.font.name = "Arial"
    nr._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    nr._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    nr.font.size = Pt(8)

    if args.core:
        summary = doc.add_paragraph()
        summary.paragraph_format.space_after = Pt(6)
        summary_run = summary.add_run(
            "Best VeriSight result on the final HaluEval QA evaluation: cross-encoder NLI claim verification + semantic evidence reranking + source-quality checks achieved 68.25% accuracy, 71.00% Macro-F1 and 72.51% hallucination-risk F1. Paper results below are not ranked against one another because their datasets and metrics differ."
        )
        summary_run.bold = True
        summary_run.font.name = "Arial"
        summary_run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        summary_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        summary_run.font.size = Pt(7.2)
        summary_run.font.color.rgb = RGBColor(31, 78, 120)

    if args.core:
        headers = ["No.", "Year", "Paper name", "Author(s)", "Publication", "Drawback in current systems", "Solution proposed", "Relevance to VeriSight", "Method used", "Reported result / best finding"]
        widths = [0.22, 0.30, 1.18, 0.66, 0.60, 1.30, 1.30, 1.28, 1.48, 1.65]
        table_font_size = 5.8
    else:
        headers = ["No.", "Year", "Paper name and direct link", "Author(s)", "Publication", "Drawback in current systems", "Solution proposed", "Relevance to VeriSight"]
        widths = [0.28, 0.38, 1.68, 0.82, 0.8, 2.05, 2.05, 2.34]
        table_font_size = 7.5
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cells = table.rows[0].cells
    for cell, header, width in zip(header_cells, headers, widths):
        set_width(cell, width)
        set_cell_margins(cell)
        shade(cell, "D9EAF7")
        write_cell(cell, header, font_size=table_font_size, bold=True, color=(31, 78, 120))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_repeat_table_header(table.rows[0])

    for index, (year, title_text, authors, venue, drawback, solution, relevance, url) in enumerate(papers, start=1):
        row = table.add_row()
        cells = row.cells
        method, result = CORE_METHOD_AND_RESULT.get(title_text, ("", ""))
        values = (
            [str(index), year, "", authors, venue, drawback, solution, relevance, method, result]
            if args.core
            else [str(index), year, "", authors, venue, drawback, solution, relevance]
        )
        for cell, value, width in zip(cells, values, widths):
            set_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if value:
                write_cell(cell, value, font_size=table_font_size)
        paper_p = cells[2].paragraphs[0]
        paper_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paper_p.paragraph_format.space_after = Pt(1)
        add_hyperlink(paper_p, title_text, url, font_size=table_font_size, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("VeriSight literature survey | Official publisher links were checked on 11 August 2026")
    footer_run.font.name = "Arial"
    footer_run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    footer_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(89, 89, 89)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
